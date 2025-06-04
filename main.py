import os
import sys
import subprocess
import subprocess
import sys
import json

#========================Environment Variables======================================
def set_env():
    keys_path = 'third_party_keys.json'
    with open(keys_path) as f:
        keys = json.load(f)

    os.environ['USER_AGENT'] = 'myagent'
    os.environ["LANGSMITH_TRACING"] = "true"
    os.environ["LANGSMITH_API_KEY"] = keys["LANGSMITH_API_KEY"]
    os.environ["GROQ_API_KEY"] = keys["GROQ_API_KEY"]
    os.environ["HF_TOKEN"] = keys["HF_TOKEN"]
    os.environ["MISTRAL_API_KEY"] = keys["MISTRAL_API_KEY"]

set_env()


#========================Packages installation======================================
def install(*packages):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", *packages])
    except subprocess.CalledProcessError as e:
        print(f"Error installing packages: {e}")
        sys.exit(1)

def validate_and_install(packages):
    for package in packages:
        try:
            __import__(package)
            print(f"{package} is already installed.")
        except ImportError:
            print(f"{package} not found. Installing...")
            install(package)

packages_to_validate = [
    "langchain_groq",
    "langchain_mistralai",
    "langchain_community",
    "beautifulsoup4",
    "langgraph",
    "pandas",
    "python-docx"
]

validate_and_install(packages_to_validate)


#========================Libraries=================================================
from langchain.chat_models import init_chat_model
from langchain_mistralai import MistralAIEmbeddings
from langchain_community.document_loaders import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.vectorstores import InMemoryVectorStore
from langgraph.graph import MessagesState, StateGraph
from langchain_core.tools import tool
from langchain_core.messages import SystemMessage
from langgraph.prebuilt import ToolNode
from langgraph.graph import END
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import MemorySaver
from langchain.schema import Document
from docx import Document as DocxDocument
import bs4


#========================Model and Database Initialization=========================
llm = init_chat_model("llama3-8b-8192", model_provider="groq")
embeddings = MistralAIEmbeddings(model="mistral-embed")
vector_store = InMemoryVectorStore(embeddings)


#========================Read Manuals and Store Data===============================
manuals_folder = "manuals"
file_path = os.path.join(manuals_folder, "PAAIProject.docx")

if os.path.exists(file_path):
    doc = DocxDocument(file_path)
    file_content = "\n".join([para.text for para in doc.paragraphs])
else:
    raise FileNotFoundError(f"The file {file_path} does not exist.")

docs = [Document(page_content=file_content)]
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
all_splits = text_splitter.split_documents(docs)
response = vector_store.add_documents(documents=all_splits)
graph_builder = StateGraph(MessagesState)


#========================Tools===========================================
@tool(response_format="content_and_artifact")
def retrieve(query: str):
    """Retrieve information related to a query."""
    retrieved_docs = vector_store.similarity_search(query, k=2)
    serialized = "\n\n".join(
        (f"Source: {doc.metadata}\n" f"Content: {doc.page_content}")
        for doc in retrieved_docs
    )
    return serialized, retrieved_docs


#========================Resonse Generation===============================
def query_or_respond(state: MessagesState):
    system_message = SystemMessage(
        "You are an assistant specializing in Google Merchant Center." \
        "Use the 'retrieve' tool only for questions about setting up and managing Google Merchant Center accounts, including topics like" \
        "account creation, business verification, shipping, sales tax, product listing requirements, and Google policies." \
        "For greetings or general questions unrelated to Google Merchant Center, provide a direct answer without calling any tools."
    )

    messages = [system_message] + state["messages"]
    llm_with_tools = llm.bind_tools([retrieve])
    response = llm_with_tools.invoke(messages)
    tool_calls = getattr(response, "tool_calls", None)

    if tool_calls is not None and not tool_calls:
        fallback = {"role": "ai", "content": "Hello! How can I assist you today?"}
        return {"messages": [fallback]}
    return {"messages": [response]}


def generate(state: MessagesState):
    recent_tool_messages = [] # Get generated ToolMessages

    for message in reversed(state["messages"]):
        if message.type == "tool":
            recent_tool_messages.append(message)
        else:
            break

    tool_messages = recent_tool_messages[::-1]

    docs_content = "\n\n".join(tool_message.content for tool_message in tool_messages if hasattr(tool_message, 'content'))

    system_message_content = (
        "You are an assistant for question-answering tasks. "
        "Use the following pieces of retrieved context to answer "
        "the question. If you don't know the answer, say that you "
        "don't know. Use three sentences maximum and keep the "
        "answer concise."
        "\n\n"
        f"{docs_content}"
    )

    conversation_messages = [
        message
        for message in state["messages"] if message.type in ("human", "system") or (message.type == "ai" and not message.tool_calls)
    ]

    prompt = [SystemMessage(system_message_content)] + conversation_messages
    response = llm.invoke(prompt) # Here we start the generation, see the graph in the sections below
    return {"messages": [response]}


#========================Graph and Execution Setup====================
tools = ToolNode([retrieve])
graph_builder.add_node(query_or_respond)
graph_builder.add_node(tools)
graph_builder.add_node(generate)
graph_builder.set_entry_point("query_or_respond")

graph_builder.add_conditional_edges(
    "query_or_respond",
    tools_condition,
    {END: END, "tools": "tools"},
)

graph_builder.add_edge("tools", "generate")
graph_builder.add_edge("generate", END)

memory = MemorySaver()
graph = graph_builder.compile(checkpointer=memory)
config = {"configurable": {"thread_id": "test_123"}}


#========================Fronetend===============================
if __name__ == "__main__":
    
    # Prompt testing
    # input_message = "Hey GPT, what's up!"
    # for step in graph.stream(
    #     {"messages": [{"role": "user", "content": input_message}]},
    #     stream_mode="values",
    #     config=config
    # ):
    #     step["messages"][-1].pretty_print()

    # UI Code
    from flask import Flask, render_template, request

    def get_response(input_message: str):
        final_response = None
        for step in graph.stream(
            {"messages": [{"role": "user", "content": input_message}]},
            stream_mode="values",
            config={"configurable": {"thread_id": "test_123"}}
        ):
            last_message = step["messages"][-1].content
            if last_message:
                final_response = last_message
        return final_response if final_response else "No response"
    
    app = Flask(__name__)

    @app.route("/", methods=["GET", "POST"])
    def index():
        response = None
        if request.method == "POST":
            input_message = request.form["input_message"]
            response = get_response(input_message) # Fetch the response from the graph
        return render_template("index.html", response=response)

    if __name__ == "__main__":
        app.run(debug=True, host='0.0.0.0', port=5001)